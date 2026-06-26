#!/usr/bin/env python3
"""Actualiza los documentos Word de Proyecto Centinela manteniendo estilo Normal."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOCS = Path(__file__).resolve().parent.parent / "Documentos"


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def insert_block_after(paragraph: Paragraph, lines: list[str]) -> Paragraph:
    current = paragraph
    for line in lines:
        current = insert_paragraph_after(current, line)
    return current


def find_paragraph(doc: Document, contains: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if contains in p.text:
            return p
    return None


def update_requirements_doc() -> None:
    path = DOCS / "ProyectoPersonalSeguridad.docx"
    doc = Document(path)

    # Título oficial
    if doc.paragraphs:
        doc.paragraphs[0].text = (
            'Análisis de requerimientos — Proyecto Centinela (Sistema de alerta comunitaria)'
        )

    stack = find_paragraph(doc, "Stack Tecnológico")
    if stack:
        insert_block_after(
            stack,
            [
                "",
                "Política de Infraestructura: Proyecto Supabase Dedicado (Aislamiento de RECI)",
                "Regla obligatoria: Centinela utilizará un proyecto Supabase completamente nuevo e independiente del proyecto RECI. Nunca compartir el mismo project ref, base de datos, Storage bucket ni claves API entre ambos sistemas.",
                "Nombre sugerido del proyecto en Supabase: centinela-mvp (o centinela-pilot).",
                "Variables de entorno separadas: CENTINELA_SUPABASE_URL y CENTINELA_SUPABASE_ANON_KEY. No reutilizar archivos .env del proyecto RECI.",
                "Organización: si ambos proyectos viven en la misma cuenta de Supabase, usar carpetas/proyectos distintos en el dashboard; los cambios de esquema, RLS o Edge Functions de Centinela nunca deben ejecutarse contra la instancia de RECI.",
                "",
                "Estrategia de Despliegue por Plataforma (Decisión v1.0)",
                "Piloto principal: Android (APK/AAB vía Google Play Console o distribución interna).",
                "Validación secundaria: iOS mediante build Flutter en dispositivo de prueba (TestFlight opcional en fase Beta). Flutter compila ambas plataformas desde el mismo código; el piloto Alfa en Jipijapa prioriza Android por alcance y facilidad de distribución comunitaria.",
            ],
        )

    doc.save(path)


def update_sdd() -> None:
    path = DOCS / "Documento de Diseño del Sistema (SDD) - Proyecto Centinela.docx"
    doc = Document(path)

    score = find_paragraph(doc, "score_confiabilidad")
    if score:
        insert_block_after(
            score,
            [
                "\t•\tultima_ubicacion (Point / Geometría Espacial): Última coordenada GPS reportada por el dispositivo. Requerida para el motor de geofencing (consulta de usuarios dentro del radio).",
                "\t•\tubicacion_actualizada_en (Timestamp): Momento de la última actualización de ubicación. Permite descartar tokens obsoletos en consultas geoespaciales.",
            ],
        )

    wireframes = find_paragraph(doc, "Diseño de Interfaz y Estructura Visual (Wireframes del MVP)")
    if wireframes:
        insert_block_after(
            wireframes,
            [
                "Estado del diseño en Figma (Wireframes MVP): Pantallas 1–3 completadas (Home, Emisión, Detalle receptor). Design system base definido (colores, Inter).",
                "Pantallas pendientes de diseño (prioridad antes de Sprint 2):",
                "\t•\tPantalla 0 — Onboarding y permisos (GPS + notificaciones push).",
                "\t•\tPantalla 0b — Login / Registro (Google, Apple, teléfono o email).",
                "\t•\tPantalla 3b — Detalle de alerta (vista Emisor): botones Marcar como Resuelto y Reportar falsa alarma.",
                "\t•\tPantalla 3c — Confirmación del flujo Lo vi (envío de coordenadas).",
                "\t•\tPantalla Web — Deep Link público (alerta activa y caso resuelto con Open Graph).",
                "\t•\tPantalla Legal — Términos y condiciones / consentimiento LOPDP.",
            ],
        )

    closing = find_paragraph(doc, "Con esto, el mensaje es un dardo")
    if closing:
        insert_block_after(
            closing,
            [
                "",
                "Notas de Infraestructura Supabase",
                "Centinela debe usar un proyecto Supabase exclusivo, separado del proyecto RECI. Ver sección correspondiente en el documento de requerimientos. El esquema SQL inicial vive en supabase/migrations/ del repositorio del proyecto.",
            ],
        )

    doc.save(path)


def update_sprint_backlog() -> None:
    path = DOCS / "Product Backlog y Plan de Sprints - Proyecto Centinela.docx"
    doc = Document(path)

    supabase_task = find_paragraph(doc, "Tarea 0.2:")
    if supabase_task:
        supabase_task.text = (
            "Tarea 0.2: Crear un proyecto Supabase NUEVO e independiente de RECI "
            "(nombre sugerido: centinela-mvp). Ejecutar los scripts SQL en supabase/migrations/ "
            "para crear las tablas usuarios, alertas_desaparecidos y reacciones_avistamientos "
            "con extensión PostGIS habilitada. Configurar RLS básico y bucket de Storage para fotos."
        )

    last = doc.paragraphs[-1]
    for p in reversed(doc.paragraphs):
        if p.text.strip():
            last = p
            break

    insert_block_after(
        last,
        [
            "Tarea 3.6: Implementar actualización periódica de ultima_ubicacion en la tabla usuarios "
            "cuando la app está activa o en segundo plano (requerido para geofencing).",
            "Tarea 3.7: Implementar post-moderación básica: botón Reportar alerta falsa, límite de "
            "alertas para cuentas nuevas y lógica de score_confiabilidad.",
            "",
            "Sprint 4: Pruebas, Legal y Piloto Alfa (El Despegue)",
            "Objetivo: validar el sistema completo en condiciones reales antes del piloto en Jipijapa.",
            "Tarea 4.1: Pruebas de estrés de notificaciones push y latencia de consultas PostGIS en radio 5 km.",
            "Tarea 4.2: Pruebas del flujo Deep Link + Open Graph en WhatsApp (imagen < 300 KB, caso resuelto).",
            "Tarea 4.3: Implementar pantallas legales (Términos y Condiciones, consentimiento LOPDP).",
            "Tarea 4.4: Corrección de bugs, optimización de batería en segundo plano y compresión de fotos.",
            "Tarea 4.5: Empaquetado Android (APK/AAB) y build de prueba iOS en dispositivo físico.",
            "Tarea 4.6: Simulacro controlado en Jipijapa (Prueba Alfa) con métricas: tiempo de emisión, "
            "latencia push, tasa de apertura, avistamientos registrados.",
            "",
            "Criterios de Aceptación por Requerimiento (MVP)",
            "RF-01: El emisor completa el reporte (foto + campos obligatorios) en menos de 20 segundos en red 4G.",
            "RF-02: Tras emitir, los dispositivos dentro del radio configurado reciben push en menos de 10 segundos.",
            "RF-03: El botón Lo vi guarda coordenadas del testigo y notifica al emisor sin exponer teléfono público.",
            "RF-04: Compartir en WhatsApp genera mensaje preformateado y tarjeta OG con foto; enlace válido sin app instalada.",
            "Post-moderación: Marcar como resuelto oculta datos en app y cambia OG a Caso resuelto.",
            "",
            "Definition of Done (DoD) — Incremento considerado terminado cuando:",
            "• Código en repositorio Git con revisión propia y sin secretos expuestos.",
            "• Funcionalidad probada en dispositivo Android físico.",
            "• Variables de entorno apuntan al proyecto Supabase centinela-mvp (nunca RECI).",
            "• Políticas RLS activas en tablas expuestas.",
            "• Documentación del sprint actualizada si hubo cambios de alcance.",
            "",
            "Sprint 5: Fixes Piloto Alfa (Post-prueba en campo)",
            "Objetivo: corregir hallazgos de la prueba con 3 Android en Jipijapa.",
            "Tarea 5.1: WhatsApp (manifest Android + intent), mapa Home usable, emisión con pin en mapa y texto de último lugar visto.",
            "Tarea 5.2: Lo vi con confirmación en mapa; Realtime en reacciones_avistamientos; resumen de avistamientos para emisor.",
            "Tarea 5.3: Firebase FCM para push con app cerrada (después de validar 5.1 en campo).",
        ],
    )

    doc.save(path)


if __name__ == "__main__":
    update_requirements_doc()
    update_sdd()
    update_sprint_backlog()
    print("Documentos actualizados correctamente.")
