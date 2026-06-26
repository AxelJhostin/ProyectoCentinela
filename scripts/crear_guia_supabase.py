#!/usr/bin/env python3
"""Genera la Guía de Aislamiento Supabase para Centinela."""

from pathlib import Path

from docx import Document

DOCS = Path(__file__).resolve().parent.parent / "Documentos"
OUTPUT = DOCS / "Guia-Aislamiento-Supabase-Centinela.docx"


def main() -> None:
    doc = Document()
    doc.add_paragraph("Guía de Aislamiento Supabase — Proyecto Centinela")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Propósito: garantizar que el desarrollo de Centinela no afecte al proyecto RECI "
        "ni a ningún otro sistema que comparta la misma cuenta de Supabase."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Regla de oro")
    doc.add_paragraph(
        "Centinela = proyecto Supabase nuevo. RECI = proyecto Supabase existente. "
        "Nunca mezclar URLs, anon keys, service role keys, buckets ni migraciones."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Pasos para crear el entorno Centinela (Sprint 0.2)")
    steps = [
        "1. Ir a supabase.com/dashboard y crear proyecto nuevo.",
        "2. Nombre sugerido: centinela-mvp (región cercana a Ecuador si está disponible).",
        "3. Guardar URL y anon key en un archivo .env.local SOLO dentro de ProyectoEmilia/.",
        "4. Prefijo de variables: CENTINELA_SUPABASE_URL, CENTINELA_SUPABASE_ANON_KEY.",
        "5. Ejecutar supabase/migrations/20250625000000_initial_schema.sql en el SQL Editor del proyecto nuevo.",
        "6. Crear bucket Storage: centinela-fotos (privado; acceso vía signed URLs en Sprint 2).",
        "7. Verificar en el dashboard que el project ref NO coincide con el de RECI.",
    ]
    for s in steps:
        doc.add_paragraph(s)
    doc.add_paragraph("")
    doc.add_paragraph("Qué NO hacer")
    donts = [
        "• No copiar .env de RECI a este repositorio.",
        "• No ejecutar migraciones de Centinela en la base de datos de RECI.",
        "• No reutilizar nombres de tablas conflictivos si algún día compartieran org (Centinela usa usuarios, alertas_desaparecidos, reacciones_avistamientos en su propia BD).",
        "• No commitear claves API al repositorio Git.",
        "• No vincular el MCP de Supabase en Cursor al proyecto RECI cuando trabajes en Centinela.",
    ]
    for d in donts:
        doc.add_paragraph(d)
    doc.add_paragraph("")
    doc.add_paragraph("Trabajo con Cursor / MCP Supabase")
    doc.add_paragraph(
        "Al usar el plugin Supabase en Cursor, selecciona explícitamente el proyecto centinela-mvp "
        "antes de ejecutar SQL o revisar tablas. Si el MCP queda conectado a RECI, desconéctalo o "
        "cambia de proyecto antes de cualquier operación de Centinela."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Checklist antes de cada sesión de desarrollo")
    checklist = [
        "□ ¿Estoy en la carpeta ProyectoEmilia?",
        "□ ¿Mis variables de entorno apuntan a centinela-mvp?",
        "□ ¿El SQL que voy a ejecutar es de supabase/migrations/ de Centinela?",
        "□ ¿No hay claves de RECI en mi .env actual?",
    ]
    for c in checklist:
        doc.add_paragraph(c)
    doc.add_paragraph("")
    doc.add_paragraph("Recuperación si hubo confusión")
    doc.add_paragraph(
        "Si accidentalmente ejecutaste SQL de Centinela en RECI: detén inmediatamente, "
        "documenta qué script corriste, revisa el historial del SQL Editor de Supabase y "
        "restaura desde backup de RECI si es necesario. Por eso esta guía exige proyecto separado."
    )
    doc.save(OUTPUT)
    print(f"Creado: {OUTPUT}")


if __name__ == "__main__":
    main()
